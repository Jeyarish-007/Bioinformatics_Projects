import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# =============== CONFIGURATION ===============

excel_path = r"D:\Shodhika\326_rows.xlsx"         # <-- Path to your Excel file
output_docx_path = r"D:\Shodhika\326_Case_Records.docx"  # <-- Path for output Word file

# =============== LOAD DATA ===============
data = pd.read_excel(excel_path)
data.columns = data.columns.str.strip()  # Clean column names, remove spaces

fields = [
    'Case No', 'Age', 'Gender',
    'Diagnosis', 'Relapse', 'Date of sample collection'
]

# =============== CREATE WORD DOC ===============
doc = Document()

for idx, row in data.iterrows():
    # Add centered title
    title_paragraph = doc.add_paragraph("Case Report Form")
    title_paragraph.style = doc.styles['Heading 1']
    title_paragraph.alignment = 1  # 0=left, 1=center

    # Add 5 blank lines after the title
    for _ in range(5):
        doc.add_paragraph()

    # Create and fill the table
    table = doc.add_table(rows=len(fields), cols=2)
    table.autofit = False

    for i, field in enumerate(fields):
        # Left cell: Field name
        cell_left = table.cell(i, 0)
        cell_left.text = field
        for paragraph in cell_left.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

        # Right cell: Value
        val = row[field] if field in row else ''

        # Format date field if needed
        if field == 'Date of sample collection':
            if pd.notna(val):
                # Format if it's a date or can be parsed as a date
                if isinstance(val, pd.Timestamp):
                    val = val.strftime('%d-%m-%Y')
                else:
                    try:
                        val = pd.to_datetime(val).strftime('%d-%m-%Y')
                    except Exception:
                        val = str(val)
            else:
                val = ''
        else:
            val = str(val) if pd.notna(val) else ''

        cell_right = table.cell(i, 1)
        cell_right.text = val
        for paragraph in cell_right.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

    # Add table borders using XML
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Remove any existing borders
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    borders = parse_xml(r'''
        <w:tblBorders %s>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
    ''' % nsdecls('w'))
    tblPr.append(borders)

    # Page break after each record, except the last
    if idx < len(data) - 1:
        doc.add_page_break()

# =============== SAVE DOC ===============
doc.save(output_docx_path)
print(f"Document created successfully: {output_docx_path}")
